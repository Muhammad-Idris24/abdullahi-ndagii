from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/abdullahi-ndagi-adamu-what-we-carry-kafart-6-proposal.docx")
INK = RGBColor(26, 25, 23)
ASH = RGBColor(102, 97, 92)
BONE = "F4EFE8"
BORDER = "D9D3CB"


def font(run, name="Cormorant Garamond", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_paragraph(doc, text="", style="body", align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if style == "body":
        fmt.space_after = Pt(8)
        fmt.line_spacing = 1.33
        run = p.add_run(text)
        font(run, name="Garamond", size=11)
    elif style == "h1":
        fmt.space_before, fmt.space_after = Pt(18), Pt(7)
        run = p.add_run(text)
        font(run, size=16, bold=True)
    elif style == "h2":
        fmt.space_before, fmt.space_after = Pt(11), Pt(5)
        run = p.add_run(text)
        font(run, size=12.5, bold=True)
    elif style == "kicker":
        fmt.space_after = Pt(4)
        run = p.add_run(text.upper())
        font(run, name="Arial", size=8.5, color=ASH, bold=True)
        run.font.letter_spacing = Pt(1.3) if hasattr(run.font, "letter_spacing") else None
    return p


def add_meta_table(doc):
    data = [
        ("Artist", "Abdullahi Ndagi Adamu"),
        ("Medium", "Acrylic and spray paint on canvas"),
        ("Collection", "3 works"),
        ("Dimensions", "3 x 4 ft (36 x 48 in) each"),
        ("Cultural context", "Kupa community, Nupe culture, Kogi State"),
        ("Primary source", "Personal photographs, memories and lived experience"),
    ]
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, pair in zip(table.rows, [(data[0], data[1]), (data[2], data[3]), (data[4], data[5])]):
        for cell, (label, value) in zip(row.cells, pair):
            cell.width = Inches(3.25)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade(cell, BONE)
            cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(label.upper())
            font(r, name="Arial", size=7.5, color=ASH, bold=True)
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            font(r, name="Garamond", size=10.5, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_work(doc, title, text):
    add_paragraph(doc, title, "h2")
    add_paragraph(doc, text)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("ABDULLAHI NDAGI ADAMU  |  WHAT WE CARRY  |  KAFART 6 PROPOSAL")
    font(r, name="Arial", size=7.5, color=ASH, bold=True)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_footer(section)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("PROPOSAL FOR KAFART 6")
    font(r, name="Arial", size=9, color=ASH, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("What We Carry")
    font(r, size=30, color=INK, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("A three-work painting collection on food, memory and cultural continuity")
    font(r, name="Garamond", size=13, color=ASH, italic=True)

    add_meta_table(doc)

    add_paragraph(doc, "Project Proposal", "h1")
    add_paragraph(doc, "What We Carry is a three-part painting series rooted in my experience growing up within the Kupa community, part of the wider Nupe cultural heritage of Kogi State. The work begins with something ordinary: smoked fish.")
    add_paragraph(doc, "Growing up in Lokoja, I remember relatives travelling between our ancestral communities and the city with smoked fish to share. When I was away at school, my father would sometimes send it with me for my principal, guardian and relatives. I did not fully understand the gesture then; looking back, I recognise that I was carrying a small piece of home. Giving the fish extended familiarity, affection and belonging across distance.")
    add_paragraph(doc, "The series became more personal as I began documenting my mother, the Queen Mother of Kupa Kingdom, participating in the same everyday practice. In one photograph, our hands meet as we break a smoked fish together. In another, her hands arrange fish to be shared with members of her community. These domestic moments hold histories of movement, family, care and memory. The third work will draw from new photographic documentation of how fish is preserved and stored within the community.")

    add_paragraph(doc, "Resonance with the Exhibition Theme", "h1")
    add_paragraph(doc, "Claypots: Food, Body and Memory offers an opportunity to consider food beyond nourishment. Like the claypot, smoked fish becomes a vessel: it holds the taste of home, the labour of preservation, and the generosity that travels from one person to another. In a time of industrial food systems, this riverside practice persists as a quiet act of cultural continuity.")
    add_paragraph(doc, "Through painting, the project explores food as a vessel of memory, love, identity and belonging. My mother carries a tradition; I carry the memory of witnessing it. The series asks what we take with us when we leave home, what we offer others as a way of sharing where we come from, and what we carry forward for those who come after us.")

    add_paragraph(doc, "Proposed Works", "h1")
    add_work(doc, "The Hands That Carry", "An intimate moment between mother and son as they break apart smoked fish from the River Niger. The work focuses on hands rather than faces, framing touch as a form of cultural transmission. My mother's hands represent knowledge inherited through lived practice; my own hands represent the act of witnessing and carrying that knowledge forward. The painting considers what is passed between generations through gestures that are rarely recorded but continuously repeated.")
    add_work(doc, "The Queen's Gift", "A portrait of the Queen Mother of Kupa Kingdom preparing smoked fish to share with members of her community. The work considers giving not simply as generosity, but as cultural responsibility. Fish moves from the river through the hands of women who prepare and preserve it, and finally into the hands of another person. Food becomes a carrier of relationship and belonging, reflecting my mother's role within both family and community.")
    add_work(doc, "A Piece of Home", "A reflection on the journey of smoked fish away from its place of origin and into the hands of someone living elsewhere. Inspired by childhood experiences of receiving and carrying smoked fish while away from home, this work considers how food collapses physical distance. A simple gift becomes a reminder of family, landscape and belonging; the River Niger remains present as an origin held within the object itself.")

    add_paragraph(doc, "Production Plan", "h1")
    for item in [
        "Medium: Acrylic and spray paint on canvas",
        "Number of works: 3",
        "Dimensions: 3 x 4 ft (36 x 48 in) each",
        "Production period: August-October 2026",
        "Research: Personal photographs, family memories and Kupa cultural knowledge",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("• ")
        font(r, name="Arial", size=10, color=ASH)
        r = p.add_run(item)
        font(r, name="Garamond", size=11)

    add_paragraph(doc, "Artist Statement", "h1")
    add_paragraph(doc, "Abdullahi Ndagi Adamu is a Nigerian visual artist from Lokoja, Kogi State, and a member of the Kupa community within the wider Nupe cultural heritage. His practice is rooted in personal experiences, memory, presence, and the ordinary moments that carry deeper meaning.")
    add_paragraph(doc, "Working primarily through painting, Abdullahi transforms personal photographs, lived experiences, family histories, and cultural observations into visual stories. His work looks closely at moments that might otherwise be overlooked, exploring how they hold questions of identity, belonging, place, memory, and human connection.")
    add_paragraph(doc, "His practice is particularly interested in the relationship between people and their environment, and in how everyday experiences can become vessels for cultural memory. Through painting, he documents moments he feels are worth remembering, while reflecting on what individuals inherit, carry, preserve, and pass forward.")

    OUTPUT.parent.mkdir(exist_ok=True)
    doc.core_properties.title = "What We Carry - KAFART 6 Proposal"
    doc.core_properties.author = "Abdullahi Ndagi Adamu"
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
